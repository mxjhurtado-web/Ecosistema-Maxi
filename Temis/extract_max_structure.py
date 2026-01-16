#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Extract content from Proyecto Max.docx to analyze structure
"""

from docx import Document
import json

doc = Document('Proyecto Max.docx')

# Extract all paragraphs with their styles
content = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        content.append({
            'index': i,
            'text': para.text[:200],  # First 200 chars
            'style': para.style.name if para.style else 'Normal'
        })

# Print first 100 paragraphs
for item in content[:100]:
    print(f"{item['index']:3d} | {item['style']:20s} | {item['text']}")

print(f"\n\nTotal paragraphs: {len(content)}")
print(f"Total paragraphs in document: {len(doc.paragraphs)}")
