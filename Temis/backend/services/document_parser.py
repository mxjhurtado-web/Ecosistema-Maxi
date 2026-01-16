#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Document Parser Service for TEMIS
Extracts text from various document formats
"""

from typing import Optional
import io


class DocumentParser:
    """Parse documents and extract text"""
    
    @staticmethod
    def extract_text(file_content: bytes, file_extension: str) -> Optional[str]:
        """
        Extract text from document
        Args:
            file_content: Binary content of file
            file_extension: File extension (.docx, .pdf, .txt)
        Returns:
            Extracted text or None if error
        """
        try:
            if file_extension.lower() == '.docx':
                return DocumentParser._extract_from_docx(file_content)
            elif file_extension.lower() == '.pdf':
                return DocumentParser._extract_from_pdf(file_content)
            elif file_extension.lower() in ['.txt', '.md']:
                return file_content.decode('utf-8')
            else:
                # Try as text
                return file_content.decode('utf-8')
        except Exception as e:
            print(f"Error extracting text: {e}")
            return None
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            raise
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            raise
